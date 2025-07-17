import pdfplumber
from docx import Document


def extract_text_from_pdf(file_stream):
    text = ""
    with pdfplumber.open(file_stream) as pdf:
        for page in pdf.pages:
            text += page.extract_text() or ""
    return text.strip()

def extract_text_from_docx(file_like_obj):
    doc = Document(file_like_obj)
    lines = []

    # Extract normal paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text.strip())

    # Extract tables (row by row)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                lines.append(row_text)

    return "\n".join(lines)

def extract_text_from_txt(file_like_obj):
    return file_like_obj.read().decode("utf-8", errors="ignore")
